"""
Import Teotitlan 2025: CSV + DOCX -> SQLite (backend/instance/portal.db)
Idempotent: upsert by numero_control, handle generic emails, hash passwords,
create canonical carreras/materias, import calificaciones.
Verify mail mock.
"""
import os
import sys
import csv
import re
import shutil
import secrets
import string
import unicodedata
from datetime import datetime
from pathlib import Path
import glob

# --- Paths ---
BASE_DIR = Path(__file__).resolve().parent
DB_PATH = BASE_DIR / "instance" / "portal.db"
CSV_PATH = Path(r"C:\Users\Dario\Desktop\portal de alumnos\BOLETAS TEOTITLAN 2025\Formulario sin título.csv")
BOLETAS_ROOT = Path(r"C:\Users\Dario\Desktop\portal de alumnos\BOLETAS TEOTITLAN 2025")

# --- Normalization helpers ---
def strip_accents(s: str) -> str:
    if not s:
        return ""
    return "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")

def normalize_carrera_raw(s: str) -> str:
    """Map variant -> canonical slug."""
    if not s:
        return "pedagogia"
    t = strip_accents(s.strip().lower())
    t = re.sub(r"\s+", " ", t).strip()
    # Remove prefixes
    # common patterns
    if "pedagog" in t:
        return "pedagogia"
    if "enfermer" in t:
        return "enfermeria"
    if "derecho" in t:
        return "derecho"
    if "psicolog" in t:
        return "psicologia"
    if "deporte" in t or "ciencias del deporte" in t:
        return "ciencias_del_deporte"
    if "conta" in t:  # contaduria, contabilidad
        return "contaduria"
    if "sistem" in t or "comput" in t or "isc" in t:
        return "sistemas"
    return "pedagogia"  # fallback

def title_case(s: str) -> str:
    if not s:
        return s
    s = re.sub(r"\s+", " ", s.strip())
    # Title case preserving accents? Use .title after lower
    return " ".join(w.capitalize() for w in s.split(" "))

def normalize_email(e: str) -> str:
    if not e:
        return ""
    return e.strip().lower()

def is_valid_email(e: str) -> bool:
    if not e:
        return False
    e = e.strip()
    if " " in e:
        return False
    if "@" not in e:
        return False
    if e.count("@") != 1:
        return False
    local, domain = e.split("@", 1)
    if not local or not domain:
        return False
    if "." not in domain:
        return False
    if len(e) < 6:
        return False
    # reject placeholders
    low = e.lower()
    if low in ("0", "o", "cero", "no", "na"):
        return False
    return True

def normalize_control(v: str, seq: int) -> str:
    if not v:
        return f"TEO2025{seq:04d}"
    s = v.strip().replace(" ", "").upper()
    # invalid markers
    if s in ("0", "O", "CERO", "CERO.", "CERO,", ""):
        return f"TEO2025{seq:04d}"
    if s == "CCT15PSU0173F":
        # This looks like clave centro trabajo, not matricula -> treat as synthetic
        return f"TEO2025{seq:04d}"
    # keep as is, but clean spaces and ensure alphanumeric
    # normalize: remove spaces, keep existing; if too short like "13" or "O" treat as synthetic
    if len(s) < 5:
        return f"TEO2025{seq:04d}"
    # Validate: should contain TEO or digits or letters; keep
    return s

def normalize_name_key(s: str) -> str:
    """For matching: lower, strip accents, collapse spaces."""
    if not s:
        return ""
    s = strip_accents(s.lower().strip())
    s = re.sub(r"\s+", " ", s)
    return s

def slug_materia(nombre: str) -> str:
    base = strip_accents(nombre.lower())
    base = re.sub(r"[^a-z0-9]+", "_", base).strip("_")
    base = base[:18].strip("_")
    return f"MAT-{base.upper()}" if base else "MAT-GENERICA"

def gen_password(length=8) -> str:
    alphabet = string.ascii_letters + string.digits
    return "".join(secrets.choice(alphabet) for _ in range(length))

def parse_folder_carrera(folder_name: str) -> tuple:
    """Extract year and carrera slug from folder name like '2024 Pedagogia 5o cuatrimestre'."""
    name = folder_name.strip()
    year = None
    m = re.match(r"^(20\d{2})", name)
    if m:
        year = int(m.group(1))
    slug = normalize_carrera_raw(name)
    # extract cuatrimestre if present
    cuatri = ""
    low = strip_accents(name.lower())
    # look for like 1er, 2o, 5o, octavo, etc.
    cuatri_map = {
        "primer": "Primero", "1er": "Primero", "1o": "Primero", "primero": "Primero",
        "segundo": "Segundo", "2o": "Segundo", "2do": "Segundo",
        "tercero": "Tercero", "3o": "Tercero",
        "cuarto": "Cuarto", "4o": "Cuarto",
        "quinto": "Quinto", "5o": "Quinto",
        "sexto": "Sexto", "6o": "Sexto",
        "septimo": "Septimo", "7o": "Septimo",
        "octavo": "Octavo", "8o": "Octavo",
        "noveno": "Noveno", "9o": "Noveno",
    }
    for k, v in cuatri_map.items():
        if k in low:
            cuatri = v
            break
    return year, slug, cuatri

def extract_docx_data(docx_path: Path):
    """Extract matricula, nombre, carrera in doc, cuatrimestre, materias."""
    from docx import Document
    try:
        doc = Document(str(docx_path))
    except Exception as e:
        return {"error": str(e), "materias": []}
    # T0: header info
    matricula = ""
    nombre_doc = ""
    carrera_doc = ""
    cuatrimestre_doc = ""
    curp = ""
    exp_date = ""
    if doc.tables:
        t0 = doc.tables[0]
        for row in t0.rows:
            txts = [c.text.strip() for c in row.cells]
            joined = " | ".join(txts)
            upper = strip_accents(joined.upper())
            if "MATRICULA" in upper:
                # parse MATRICULA: XXXXX
                for txt in txts:
                    if "MATRICULA" in strip_accents(txt.upper()):
                        # after colon
                        parts = txt.split(":", 1)
                        if len(parts) == 2:
                            matricula = parts[1].strip().replace(" ", "").upper()
                        else:
                            matricula = ""
                        break
            if "NOMBRE" in upper:
                for txt in txts:
                    if "NOMBRE" in strip_accents(txt.upper()):
                        parts = txt.split(":", 1)
                        if len(parts) == 2:
                            nombre_doc = parts[1].strip()
                        break
            if "CARRERA" in upper:
                for txt in txts:
                    if "CARRERA" in strip_accents(txt.upper()):
                        parts = txt.split(":", 1)
                        if len(parts) == 2:
                            carrera_doc = parts[1].strip()
                        # CURP also in same row second cell
                        break
            if "CUATRIMESTRE" in upper:
                for txt in txts:
                    if "CUATRIMESTRE" in strip_accents(txt.upper()):
                        parts = txt.split(":", 1)
                        if len(parts) == 2:
                            cuatrimestre_doc = parts[1].strip()
                        break
            if "FECHA" in upper:
                exp_date = joined
    # T1: materias
    materias = []
    if len(doc.tables) >= 2:
        t1 = doc.tables[1]
        # Detect header rows (first 2 rows are headers)
        for ri, row in enumerate(t1.rows):
            if ri < 2:
                continue
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 4:
                continue
            # columns: 0 No REGISTRO UNIDAD, 1 No REGISTRO PERIODO, 2 ASIGNATURAS, 3 NUMERO, 4 LETRA, 5 OBS
            asign = cells[2].strip()
            calif_str = cells[3].strip()
            if not asign:
                continue
            # Skip empty or total rows
            if asign.upper().startswith("ASIGNATURAS"):
                continue
            # calif should be numeric
            try:
                calif = float(calif_str) if calif_str else None
            except:
                calif = None
            if calif is None:
                # sometimes letter column has number?
                # try cell 4
                try:
                    calif = float(cells[4].strip()) if cells[4].strip().isdigit() else None
                except:
                    calif = None
            if calif is None and not calif_str:
                continue
            # Filter non-subject rows: if asign is numeric or empty
            if asign and calif is not None:
                materias.append({"nombre": asign.strip(), "calificacion": calif})
            elif asign and calif_str and calif_str.replace(".","").isdigit():
                try:
                    materias.append({"nombre": asign.strip(), "calificacion": float(calif_str)})
                except:
                    pass
    return {
        "matricula": matricula,
        "nombre_doc": nombre_doc,
        "carrera_doc": carrera_doc,
        "cuatrimestre_doc": cuatrimestre_doc,
        "materias": materias,
        "error": None,
    }

def main():
    import sys
    sys.path.insert(0, str(BASE_DIR))
    from app import create_app
    from models import db, Carrera, Materia, Alumno, Calificacion
    from werkzeug.security import generate_password_hash

    app = create_app()
    with app.app_context():
        # 1. Backup
        timestamp = datetime.now().strftime("%Y%m%d-%H%M")
        backup_path = DB_PATH.parent / f"portal.db.bak.{timestamp}"
        if DB_PATH.exists():
            shutil.copy2(DB_PATH, backup_path)
            print(f"[BACKUP] {backup_path}")
        else:
            print("[BACKUP] no db found, skipping")

        # 2. Ensure canonical carreras
        canon = {
            "pedagogia": ("Pedagogia", "PED"),
            "enfermeria": ("Enfermeria", "ENF"),
            "derecho": ("Derecho", "DER"),
            "psicologia": ("Psicologia", "PSI"),
            "ciencias_del_deporte": ("Ciencias del Deporte", "CDE"),
            "contaduria": ("Contaduria", "CON"),
            "sistemas": ("Sistemas", "SIS"),
        }
        carrera_map = {}  # slug -> Carrera
        for slug, (nombre, codigo) in canon.items():
            c = Carrera.query.filter_by(codigo=codigo).first()
            if not c:
                # try by nombre ilike
                c = Carrera.query.filter(db.func.lower(Carrera.nombre) == nombre.lower()).first()
                if c:
                    # keep existing but also ensure code mapping; create alias
                    # We'll create new with canonical code if not exists
                    pass
                if not c or c.codigo != codigo:
                    # create if missing code
                    exists_code = Carrera.query.filter_by(codigo=codigo).first()
                    if not exists_code:
                        c_new = Carrera(nombre=nombre, codigo=codigo, descripcion=f"Carrera {nombre} - import Teotitlan 2025", activa=True)
                        db.session.add(c_new)
                        db.session.flush()
                        c = c_new
                        print(f"[CARRERA] created {codigo} - {nombre}")
                    else:
                        c = exists_code
            carrera_map[slug] = c
        db.session.commit()
        # Also build lookup for existing carreras by strip_accents lower
        # Refresh ids
        for slug in list(carrera_map.keys()):
            carrera_map[slug] = Carrera.query.filter_by(codigo=canon[slug][1]).first()

        print(f"[CARRERAS] canonical ensured: {[(k, v.codigo, v.id) for k,v in carrera_map.items()]}")
        # For legacy compat, also map variants to canonical via normalize_carrera_raw -> slug
        # Already done.

        # 3. Parse CSV
        csv_rows = []
        if not CSV_PATH.exists():
            print(f"[CSV] not found {CSV_PATH}")
            sys.exit(1)
        # Handle encoding with BOM and latin variations; try utf-8-sig fallback
        content = None
        for enc in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                with open(CSV_PATH, encoding=enc) as f:
                    content = f.read()
                break
            except Exception:
                continue
        # Write temp utf-8 for csv parsing
        import io
        # Use csv with BOM handling
        with open(CSV_PATH, encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            fieldnames = reader.fieldnames
            # Normalize fieldnames strip
            norm_fields = {fn: fn.strip() for fn in fieldnames}
            rows = list(reader)
        # Normalize headers mapping
        def get_field(row, *candidates):
            for cand in candidates:
                for k,v in row.items():
                    if k.strip().lower() == cand.strip().lower():
                        return v
                    # handle accent mismatch for Numero de control header encoding
                    if strip_accents(k.strip().lower()) == strip_accents(cand.strip().lower()):
                        return v
            return ""
        total_csv = len(rows)
        print(f"[CSV] total rows {total_csv} fields {fieldnames}")

        # Inspect carrera variants stats (log)
        from collections import Counter, defaultdict
        carrera_counter = Counter()
        for r in rows:
            carrera_raw = get_field(r, "Carrera")
            slug = normalize_carrera_raw(carrera_raw)
            carrera_counter[slug] += 1
        print(f"[CSV] carrera normalized {dict(carrera_counter)}")

        # 4. Build CSV alumnos list (dedupe)
        seq = 1
        # Find max existing synthetic to avoid collision
        existing_controls = {a.numero_control for a in Alumno.query.all()}
        # Also need to generate synthetic counters not colliding
        def next_synthetic():
            nonlocal seq
            while True:
                cand = f"TEO2025{seq:04d}"
                seq += 1
                if cand not in existing_controls:
                    existing_controls.add(cand)
                    return cand

        # Pre-scan to count existing synthetic max
        # Start seq from 1 but ensure not colliding; if existing has TEO20250001, we will skip via loop
        email_seen = set()  # lower email seen in CSV processing
        email_seen_db = {a.email.lower() for a in Alumno.query.all()}
        numero_seen = set()  # to dedupe exact rows
        alumnos_csv = []  # list of dicts
        generic_reasons = []  # for final CSV
        dup_email_count = 0
        invalid_email_count = 0
        synthetic_control_count = 0

        # For duplicate detection by full normalized row: skip exact duplicates
        seen_row_keys = set()

        for idx, r in enumerate(rows):
            ap = get_field(r, "Apellido paterno ").strip()
            am = get_field(r, "Apellido materno").strip()
            nombres = get_field(r, "Nombres").strip()
            correo_raw = get_field(r, "Correo ").strip()
            control_raw = get_field(r, "Número de control").strip()
            pwd_raw = get_field(r, "Contraseña (anotela y guardela)").strip()
            carrera_raw = get_field(r, "Carrera").strip()

            # Normalized carrera slug
            carrera_slug = normalize_carrera_raw(carrera_raw)
            # Normalize names
            nombre_norm = title_case(nombres)
            ap_norm = title_case(ap) if ap else ""
            am_norm = title_case(am) if am else ""
            # Some rows have email in name field due to malformed CSV (row with sharonrubicarrera6@gmail.com as ap_materno)
            # Detect if nombres looks like email and ap fields swapped
            # Heuristic: if correo_raw is not email but ap or am contains @
            # Already normalize_email handles.

            email_norm = normalize_email(correo_raw)
            # Row key for dedup: normalized lower email + normalized names + control_raw strip
            row_key = (email_norm, normalize_name_key(f"{nombres} {ap} {am}"), control_raw.strip().lower())
            if row_key in seen_row_keys:
                continue
            seen_row_keys.add(row_key)

            # Validate email
            email_valid = is_valid_email(email_norm)
            needs_generic = False
            origin_reason = ""
            if not email_valid:
                needs_generic = True
                invalid_email_count += 1
                origin_reason = "csv_sin_correo" if not email_norm or email_norm in ("0","") else "csv_correo_invalido"
            elif email_norm.lower() in email_seen:
                needs_generic = True
                dup_email_count += 1
                origin_reason = "csv_duplicado"
            elif email_norm.lower() in email_seen_db:
                needs_generic = True
                dup_email_count += 1
                origin_reason = "csv_duplicado_db"
            else:
                origin_reason = "csv"

            # Normalize control
            control_norm = normalize_control(control_raw, seq)
            if control_norm.startswith("TEO2025"):
                # was synthetic -> increment seq already inside function; but normalize_control didn't call next_synthetic with loop? It uses seq directly.
                # Need to handle seq increment properly: if synthetic generated, we must bump seq and track
                # normalize_control uses seq passed; but we need to update seq externally
                # Check if control_raw was synthetic marker: then synthetic_control_count and update seq
                if control_raw.strip().lower() in ("0","o","cero","") or control_raw.strip() == "" or len(control_raw.strip())<5 or control_raw.strip().upper()=="CCT15PSU0173F" or control_raw.strip().upper()=="O":
                    # it used synthetic pattern TEO2025xxxx; ensure uniqueness
                    # normalize_control returned TEO2025 with current seq, need to ensure seq advanced
                    # It used seq value; now advance
                    seq += 1
                    synthetic_control_count += 1
                    existing_controls.add(control_norm)
                else:
                    # control_norm may be same as input but cleaned; ensure not colliding
                    if control_norm in existing_controls and control_norm != control_raw.strip():
                        # collision -> generate new
                        control_norm = next_synthetic()
                        synthetic_control_count += 1
                    else:
                        existing_controls.add(control_norm)
            else:
                # check duplicate control in DB? We'll handle upsert later; track seen
                if control_norm in existing_controls:
                    # keep but will be update not duplicate; still add
                    pass
                existing_controls.add(control_norm)

            # Email generic handling
            final_email = email_norm if (email_valid and not needs_generic) else ""
            if needs_generic or not email_valid:
                # generate generic: nombre.apellido.numero_control@generic...
                # Use generic.teotitlan.local per spec; also alumno.{numero_control}@teotitlan.fv.local fallback for doc-only
                base_local = f"{strip_accents(nombre_norm.split()[0].lower()) if nombre_norm else 'alumno'}.{strip_accents(ap_norm.lower()) if ap_norm else 'teo'}.{control_norm.lower()}".replace(" ", "")
                base_local = re.sub(r"[^a-z0-9._-]", "", base_local)
                final_email = f"{base_local}@generic.teotitlan.local"
                # Ensure uniqueness vs db and seen
                suffix = 1
                orig_email = final_email
                while final_email.lower() in email_seen or final_email.lower() in email_seen_db:
                    final_email = f"{base_local}.{suffix}@generic.teotitlan.local"
                    suffix += 1
                email_seen.add(final_email.lower())
                email_seen_db.add(final_email.lower())
            else:
                email_seen.add(email_norm.lower())
                email_seen_db.add(email_norm.lower())
                final_email = email_norm

            # Password handling
            pwd = pwd_raw.strip() if pwd_raw and pwd_raw.strip() not in ("0", "O", "") else ""
            if not pwd or len(pwd) < 4:
                pwd = gen_password(8)
                pwd_generated = True
            else:
                pwd_generated = False

            # Carrera assignment
            carrera_obj = carrera_map.get(carrera_slug)
            if not carrera_obj:
                carrera_obj = carrera_map["pedagogia"]

            alumnos_csv.append({
                "ap": ap_norm,
                "am": am_norm,
                "nombre": nombre_norm,
                "email": final_email,
                "email_valid_orig": email_valid,
                "control": control_norm,
                "password": pwd,
                "pwd_generated": pwd_generated,
                "carrera_slug": carrera_slug,
                "carrera_id": carrera_obj.id,
                "activo": True,
                "origen": origin_reason,
                "needs_generic": needs_generic or not email_valid,
                "raw_row": r,
            })

        print(f"[CSV] deduplicated alumnos_csv {len(alumnos_csv)} dup_email {dup_email_count} invalid {invalid_email_count} synthetic {synthetic_control_count}")

        # 5. Parse DOCX
        docx_files = []
        for root, dirs, files in os.walk(BOLETAS_ROOT):
            for file in files:
                if file.startswith("~$"):
                    continue
                if file.lower().endswith(".docx"):
                    p = Path(root) / file
                    docx_files.append(p)
        total_docx = len(docx_files)
        print(f"[DOCX] found {total_docx} files")

        # Parse each
        docx_data = []
        for p in docx_files:
            rel = p.relative_to(BOLETAS_ROOT)
            folder = p.parent.name
            # need to walk up to detect BAJA folder
            # check any part of relative path contains BAJA
            parts = [x.lower() for x in p.relative_to(BOLETAS_ROOT).parts]
            is_baja = any("baja" in part for part in parts)
            # also filename contains BAJA / PENDIENTE
            if "baja" in p.stem.lower():
                is_baja = True
            is_pendiente = "pendiente" in p.stem.lower() or "pendiente" in folder.lower()
            year, carrera_slug_folder, cuatri_folder = parse_folder_carrera(folder)
            # But if folder is BAJA or BAJAS, need parent's parent
            if folder.lower() in ("baja", "bajas", "baja temporal"):
                # use grandparent
                parent_folder = p.parent.parent.name
                year2, slug2, cuatri2 = parse_folder_carrera(parent_folder)
                if year2:
                    year = year2
                if slug2 != "pedagogia" or carrera_slug_folder == "pedagogia":
                    carrera_slug_folder = slug2
                if cuatri2 and not cuatri_folder:
                    cuatri_folder = cuatri2
            # Determine carrera folder fallback: if still generic, use parent's parent maybe
            info = extract_docx_data(p)
            # Determine final carrera: priority doc carrera_doc if parseable else folder slug
            # normalize doc carrera
            if info.get("carrera_doc"):
                slug_doc = normalize_carrera_raw(info["carrera_doc"])
                # if doc says contabilidad -> contaduria
                carrera_slug = slug_doc
            else:
                carrera_slug = carrera_slug_folder
            # matricula: from doc if available and valid else synthetic later
            matricula_raw = info.get("matricula", "").strip()
            # Some docs have empty matricula
            matricula_norm = matricula_raw.replace(" ", "").upper() if matricula_raw else ""
            if not matricula_norm or len(matricula_norm) < 5 or matricula_norm.upper() == "CCT15PSU0173F":
                matricula_norm = ""
            # nombre: try doc nombre else filename
            nombre_doc = info.get("nombre_doc", "").strip()
            if nombre_doc:
                # doc nombre is like "MARTINEZ FILIO JOCELIN" -> need split?
                # Keep as is, but also parse
                nombre_full_doc = title_case(nombre_doc)
            else:
                nombre_full_doc = ""
            filename_stem = p.stem
            # remove BAJA / PENDIENTE suffix from filename for name extraction
            clean_stem = re.sub(r"\s*(BAJA|PENDIENTE|BAJA TEMPORAL).*", "", filename_stem, flags=re.I).strip()
            # If doc nombre available, prefer it; else filename
            nombre_full = nombre_full_doc if nombre_full_doc else clean_stem
            # For matching, normalize
            # Also try to split nombre_full into nombre, ap, am heuristic: first token is ap_paterno? But doc format is "APELLIDO APELLIDO NOMBRE"
            # Keep simple: store full and later split? For Alumno we need nombre, apellido_paterno, apellido_materno.
            # We'll split full name: assume last 1-2 tokens are given names? Hard.
            # Approach: use filename as source but also keep doc.
            # For now store raw full.
            docx_data.append({
                "path": p,
                "rel": rel,
                "folder": folder,
                "year": year,
                "carrera_slug": carrera_slug,
                "carrera_slug_folder": carrera_slug_folder,
                "cuatri_folder": cuatri_folder,
                "cuatri_doc": info.get("cuatrimestre_doc", ""),
                "is_baja": is_baja,
                "is_pendiente": is_pendiente,
                "matricula": matricula_norm,
                "nombre_doc": nombre_full_doc,
                "nombre_file": clean_stem,
                "nombre_full": nombre_full,
                "materias": info.get("materias", []),
                "error": info.get("error"),
            })

        # Stats
        baja_count = sum(1 for d in docx_data if d["is_baja"])
        pendiente_count = sum(1 for d in docx_data if d["is_pendiente"])
        print(f"[DOCX] baja {baja_count} pendiente {pendiente_count} with materias {sum(1 for d in docx_data if d['materias'])} no_matricula {sum(1 for d in docx_data if not d['matricula'])}")

        # 6. Cruce CSV vs DOCX
        # Normalize names for matching
        def full_key(nombre_full):
            return normalize_name_key(nombre_full)

        # Build csv name index: map normalized full name (nombre + ap + am) and also nombre_doc style
        csv_name_map = {}
        for idx, ac in enumerate(alumnos_csv):
            key1 = normalize_name_key(f"{ac['nombre']} {ac['ap']} {ac['am']}")
            key2 = normalize_name_key(f"{ac['ap']} {ac['am']} {ac['nombre']}")
            csv_name_map[key1] = idx
            csv_name_map[key2] = idx

        matched = 0
        only_boleta = []
        # Also need to update alumnos_csv with docx matricula/materias if matched?
        # We'll keep docx-only alumnos separately
        docx_only = []
        for d in docx_data:
            kf = full_key(d["nombre_full"])
            kd = full_key(d["nombre_doc"]) if d["nombre_doc"] else ""
            # Also try splitting filename vs doc: try both
            matched_idx = None
            if kf in csv_name_map:
                matched_idx = csv_name_map[kf]
            elif kd and kd in csv_name_map:
                matched_idx = csv_name_map[kd]
            else:
                # try token overlap: check if any csv alumno's tokens overlap significantly
                # Simple: for each csv, check if all tokens of docx name appear in csv name
                tokens_doc = set(kf.split())
                best = None
                for key, idx in csv_name_map.items():
                    tokens_csv = set(key.split())
                    # intersection >= 2 and covers doc tokens
                    if len(tokens_doc & tokens_csv) >= 2 and len(tokens_doc & tokens_csv) / max(len(tokens_doc),1) >= 0.6:
                        best = idx
                        break
                if best is not None:
                    matched_idx = best

            if matched_idx is not None:
                matched += 1
                # enrich csv entry with doc info: if csv control was synthetic and doc has matricula, prefer doc matricula if valid & not synthetic
                # but we keep control as is to avoid breaking unique; if doc matricula is more realistic and csv was synthetic, we could update?
                # Let's not change control if doc has matricula and csv was synthetic? Spec says for alumnos solo en DOCX crear, but for matched we keep CSV.
                # Optionally update carrera if doc carrera more accurate? Keep CSV carrera.
                # Store materias for later calif import via matched mapping
                ac = alumnos_csv[matched_idx]
                # attach docx materias to alumno entry for calif creation
                if "docx_materias" not in ac:
                    ac["docx_materias"] = []
                ac["docx_materias"].extend(d["materias"])
                # also track doc path for audit
                if "docx_paths" not in ac:
                    ac["docx_paths"] = []
                ac["docx_paths"].append(str(d["rel"]))
                # if is_baja, mark alumno inactive
                if d["is_baja"]:
                    ac["activo"] = False
            else:
                only_boleta.append(d)

        print(f"[CRUCE] matched {matched} docx_only {len(only_boleta)} total_docx {total_docx}")

        # 7. Prepare alumnos for DB: alumnos_csv + docx_only alumnos
        # For docx_only, create alumno entries with generic email, synthetic control if needed
        generic_entries = []  # for CSV output
        # Counters for logs
        created = 0
        updated = 0
        generic_emails = 0
        califs_imported = 0
        materias_new = 0
        carreras_new = 0 # already counted
        docx_only_created = 0

        # Helper to split full name into nombre, ap, am for docx_only
        def split_full_name(full):
            # Heuristic: assume format is either "Nombre Apellido" or "Apellido Apellido Nombre" or "Apellido Nombre"
            # We have doc name like "MARTINEZ FILIO JOCELIN" => ap=MARTINEZ, am=FILIO, nombre=JOCELIN
            # File name like "Joselin Martinez Filio" => nombre=Joselin, ap=Martinez, am=Filio
            # Detect if first token looks like given name? Hard.
            # Choose: if doc name is all caps and has 3 tokens, treat as AP AP NOMBRE
            # If file name is Title Case with 3 tokens, treat as NOMBRE AP AP
            # We'll try to infer: if full is all uppercase => AP AM NOMBRE
            # else => NOMBRE AP AM
            full = re.sub(r"\s+", " ", full.strip())
            parts = full.split()
            if not parts:
                return ("Alumno", "Sin", "")
            # Check if full is mostly uppercase
            upper_count = sum(1 for c in full if c.isupper())
            lower_count = sum(1 for c in full if c.islower())
            if upper_count > lower_count and len(parts) >= 2:
                # AP AP NOMBRE -> last token(s) is nombre
                # Assume format: AP_PAT AP_MAT NOMBRE(S)
                if len(parts) == 2:
                    return (title_case(parts[1]), title_case(parts[0]), "")
                elif len(parts) == 3:
                    return (title_case(parts[2]), title_case(parts[0]), title_case(parts[1]))
                else:
                    # >3: first two are apellidos, rest is nombre
                    ap = title_case(parts[0])
                    am = title_case(parts[1])
                    nombre = title_case(" ".join(parts[2:]))
                    return (nombre, ap, am)
            else:
                # NOMBRE AP AM
                if len(parts) == 1:
                    return (title_case(parts[0]), "SinApellido", "")
                elif len(parts) == 2:
                    return (title_case(parts[0]), title_case(parts[1]), "")
                elif len(parts) == 3:
                    return (title_case(parts[0]), title_case(parts[1]), title_case(parts[2]))
                else:
                    nombre = title_case(parts[0])
                    ap = title_case(parts[1])
                    am = title_case(" ".join(parts[2:])) if len(parts)>2 else ""
                    # For case "Joselin Martinez Filio" -> nombre=Joselin, ap=Martinez, am=Filio perfect
                    # For 4 tokens: "Ana Maria Martinez Nepomuceno" -> nombre="Ana Maria", ap=Martinez, am=Nepomuceno? But our split would give nombre=Ana, ap=Maria, am=Martinez Nepomuceno -> not perfect but okay
                    # Improve: if 4 tokens and first two are given names?
                    if len(parts) == 4:
                        # Assume NOMBRE NOMBRE AP AM
                        nombre = title_case(" ".join(parts[:2]))
                        ap = title_case(parts[2])
                        am = title_case(parts[3])
                        return (nombre, ap, am)
                    return (nombre, ap, am)

        # Build DB name map for idempotent docx_only reuse
        db_name_map = {}
        for a in Alumno.query.all():
            key = normalize_name_key(a.nombre_completo)
            db_name_map[key] = a
            # also add reversed key? Ensure both orders map
            rev = normalize_name_key(f"{a.apellido_paterno} {a.apellido_materno or ''} {a.nombre}".strip())
            db_name_map[rev] = a

        # Process docx_only
        docx_only_alumnos = []
        for d in only_boleta:
            nombre, ap, am = split_full_name(d["nombre_full"])
            norm_key = normalize_name_key(d["nombre_full"])
            # Check DB reuse first for idempotency
            existing_by_name = db_name_map.get(norm_key)
            if not existing_by_name:
                # also try alternative split key
                alt_key = normalize_name_key(f"{nombre} {ap} {am}")
                existing_by_name = db_name_map.get(alt_key)
            if existing_by_name:
                ctrl = existing_by_name.numero_control
                email_g = existing_by_name.email
                # keep existing password hash, don't rotate
                pwd = None
                # ensure carrera matches doc? Keep existing carrera but update if doc carrera different?
                carrera_obj = carrera_map.get(d["carrera_slug"], carrera_map["pedagogia"])
                # Mark for update but reuse control/email
                # Need to ensure email_seen tracking
                email_seen.add(email_g.lower())
                email_seen_db.add(email_g.lower())
                existing_controls.add(ctrl)
                reuse_existing = True
            else:
                reuse_existing = False
                # Control
                ctrl = d["matricula"]
                if not ctrl or ctrl in existing_controls:
                    ctrl = next_synthetic()
                    synthetic_control_count += 1
                else:
                    existing_controls.add(ctrl)
                # Email generic per spec: alumno.{numero_control}@teotitlan.fv.local or generic
                base_local = f"{strip_accents(nombre.split()[0].lower()) if nombre else 'alumno'}.{strip_accents(ap.lower()) if ap else 'teo'}.{ctrl.lower()}"
                base_local = re.sub(r"[^a-z0-9._-]", "", base_local)
                email_g = f"{base_local}@teotitlan.fv.local"
                # ensure unique
                suffix = 1
                orig = email_g
                while email_g.lower() in email_seen or email_g.lower() in email_seen_db:
                    email_g = f"{base_local}.{suffix}@teotitlan.fv.local"
                    suffix += 1
                email_seen.add(email_g.lower())
                email_seen_db.add(email_g.lower())
                carrera_obj = carrera_map.get(d["carrera_slug"], carrera_map["pedagogia"])
                pwd = gen_password(8)
            # Register in map for next iterations within same run
            if norm_key not in db_name_map:
                # create placeholder mapping to avoid duplicate synthetic within same run
                # we will add after commit, but for now map synthetic
                pass
            entry = {
                "ap": ap,
                "am": am,
                "nombre": nombre,
                "email": email_g,
                "control": ctrl,
                "password": pwd,
                "carrera_slug": d["carrera_slug"],
                "carrera_id": carrera_obj.id,
                "activo": False if d["is_baja"] else True,
                "origen": "solo_boleta",
                "needs_generic": True,
                "docx_materias": d["materias"],
                "docx_paths": [str(d["rel"])],
                "is_baja": d["is_baja"],
            }
            docx_only_alumnos.append(entry)
            # For generic CSV
            generic_entries.append({
                "numero_control": ctrl,
                "nombre_completo": f"{nombre} {ap} {am}".strip(),
                "email_generico": email_g,
                "carrera": carrera_obj.nombre,
                "origen": "solo_boleta",
                "archivo_docx": str(d["rel"]),
            })

        # For csv alumnos with generic, also add to generic_entries
        for ac in alumnos_csv:
            if ac["needs_generic"]:
                generic_entries.append({
                    "numero_control": ac["control"],
                    "nombre_completo": f"{ac['nombre']} {ac['ap']} {ac['am']}".strip(),
                    "email_generico": ac["email"],
                    "carrera": carrera_map[ac["carrera_slug"]].nombre if ac["carrera_slug"] in carrera_map else ac["carrera_slug"],
                    "origen": ac["origen"],
                    "archivo_docx": ";".join(ac.get("docx_paths", [])),
                })

        # Combine all alumnos to import
        all_alumnos = alumnos_csv + docx_only_alumnos
        print(f"[IMPORT] total alumnos to upsert {len(all_alumnos)} (csv {len(alumnos_csv)} + docx_only {len(docx_only_alumnos)})")

        # 8. Upsert Carreras/Materias/Alumnos/Calificaciones
        # Track materias new
        for alum in all_alumnos:
            # Ensure carrera exists (already)
            carrera_id = alum["carrera_id"]
            # Upsert alumno by numero_control
            existing = Alumno.query.filter_by(numero_control=alum["control"]).first()
            # Also check email collision with different control: if email exists for different control, we already generated generic unique
            if existing:
                # update
                existing.nombre = alum["nombre"]
                existing.apellido_paterno = alum["ap"]
                existing.apellido_materno = alum["am"] if alum["am"] else None
                # email: update only if not colliding with other alumno
                if existing.email.lower() != alum["email"].lower():
                    # check if alum email exists for other alumno
                    conflict = Alumno.query.filter(db.func.lower(Alumno.email) == alum["email"].lower(), Alumno.id != existing.id).first()
                    if not conflict:
                        existing.email = alum["email"]
                    else:
                        # generate alternative
                        base = alum["email"].split("@")[0]
                        domain = alum["email"].split("@")[1]
                        suffix = 1
                        new_email = f"{base}.{suffix}@{domain}"
                        while Alumno.query.filter(db.func.lower(Alumno.email) == new_email.lower()).first():
                            suffix += 1
                            new_email = f"{base}.{suffix}@{domain}"
                        existing.email = new_email
                existing.carrera_id = carrera_id
                existing.activo = alum["activo"]
                # password: only rotate if alum has password (docx_only reuse has None)
                if alum.get("password"):
                    existing.set_password(alum["password"])
                # also set must_change_password? leave false
                db.session.flush()
                updated += 1
                alum_id = existing.id
                # for generic count
                if alum["needs_generic"]:
                    generic_emails += 1
            else:
                # Check email unique
                conflict = Alumno.query.filter(db.func.lower(Alumno.email) == alum["email"].lower()).first()
                final_email = alum["email"]
                if conflict:
                    base = alum["email"].split("@")[0]
                    domain = alum["email"].split("@")[1]
                    suffix = 1
                    new_email = f"{base}.{suffix}@{domain}"
                    while Alumno.query.filter(db.func.lower(Alumno.email) == new_email.lower()).first():
                        suffix += 1
                        new_email = f"{base}.{suffix}@{domain}"
                    final_email = new_email
                new_al = Alumno(
                    numero_control=alum["control"],
                    nombre=alum["nombre"],
                    apellido_paterno=alum["ap"],
                    apellido_materno=alum["am"] if alum["am"] else None,
                    email=final_email,
                    carrera_id=carrera_id,
                    activo=alum["activo"],
                )
                new_al.set_password(alum["password"])
                db.session.add(new_al)
                db.session.flush()
                alum_id = new_al.id
                created += 1
                if alum["needs_generic"]:
                    generic_emails += 1
                if alum["origen"] == "solo_boleta":
                    docx_only_created += 1

            # Materias / Calificaciones
            materias_list = alum.get("docx_materias", [])
            # also if alumno had no docx materias, we still may create empty
            for mat in materias_list:
                mat_nombre = mat["nombre"].strip()
                # Truncate long names? Keep as is but title case maybe
                # Normalize materia nombre: Title Case, strip
                mat_nombre_clean = re.sub(r"\s+", " ", mat_nombre).strip()
                # Some materia names are truncated like "INTRODUCCION AL ESTUDIO DEL DE" – keep
                # Ensure materia exists for this carrera
                existing_mat = Materia.query.filter_by(carrera_id=carrera_id, nombre=mat_nombre_clean).first()
                if not existing_mat:
                    # also try case-insensitive
                    existing_mat = Materia.query.filter(db.func.lower(Materia.nombre) == mat_nombre_clean.lower(), Materia.carrera_id == carrera_id).first()
                if not existing_mat:
                    code = slug_materia(mat_nombre_clean)
                    # ensure code unique within carrera? Check duplicate code in same carrera
                    # If code collision, append suffix
                    base_code = code
                    suffix = 1
                    while Materia.query.filter_by(carrera_id=carrera_id, codigo=code).first():
                        code = f"{base_code}_{suffix}"
                        suffix += 1
                    new_mat = Materia(carrera_id=carrera_id, nombre=mat_nombre_clean, codigo=code, creditos=0)
                    db.session.add(new_mat)
                    db.session.flush()
                    materias_new += 1
                    mat_id = new_mat.id
                else:
                    mat_id = existing_mat.id

                # Calificacion upsert: unique (alumno_id, materia_id, periodo, anio)
                # Derive periodo/anio from alum's docx info: use year and cuatri
                # For simplicity, use alum's first docx year/cuatri or default 2025
                # Find corresponding docx for this alum to get year/cuatri; fallback
                # Use year from carrera mapping? Use 2025 and periodo = "2025"
                # Better: get year from docx_data if available
                # We'll derive from alum's docx_paths: peek first
                anio = 2025
                periodo = "2025"
                # Try to get from first path's folder year
                try:
                    # alum may have docx_paths
                    if alum.get("docx_paths"):
                        rel = alum["docx_paths"][0]
                        # parse year from path
                        m = re.search(r"(20\d{2})", rel)
                        if m:
                            anio = int(m.group(1))
                        # cuatri from docx
                        # use d cuatri if matched? Find matching docx_data entry
                        for d in docx_data:
                            if str(d["rel"]) == rel:
                                cuatri = d.get("cuatri_doc") or d.get("cuatri_folder") or ""
                                if cuatri:
                                    periodo = f"{cuatri} {anio}"
                                else:
                                    periodo = f"{anio}"
                                break
                except:
                    pass

                calif_val = mat["calificacion"]
                # Check existing calificacion with same alumno, materia, periodo, anio
                # If materia truncated, there may be multiple materias with similar prefix – but treat exact match
                existing_cal = Calificacion.query.filter_by(alumno_id=alum_id, materia_id=mat_id, periodo=periodo, anio=anio).first()
                if existing_cal:
                    existing_cal.calificacion_final = float(calif_val)
                    # no duplicate increment
                else:
                    new_cal = Calificacion(
                        alumno_id=alum_id,
                        materia_id=mat_id,
                        calificacion_final=float(calif_val),
                        periodo=periodo,
                        anio=anio,
                    )
                    db.session.add(new_cal)
                    califs_imported += 1

        db.session.commit()
        print(f"[UPSERT] created {created} updated {updated} generic {generic_emails} materias_new {materias_new} califs_new {califs_imported} docx_only {docx_only_created}")

        # 9. Write generic CSV
        out_csv = DB_PATH.parent / "alumnos_genericos_para_contactar.csv"
        with open(out_csv, "w", encoding="utf-8-sig", newline="") as f:
            w = csv.DictWriter(f, fieldnames=["numero_control","nombre_completo","email_generico","carrera","origen","archivo_docx"])
            w.writeheader()
            for row in generic_entries:
                w.writerow(row)
        print(f"[CSV OUT] {out_csv} rows {len(generic_entries)}")

        # 10. Verify mail mock
        from utils.mail import send_credentials_email, render_credentials_email, _get_smtp_config
        import utils.mail as mail_mod
        test_html = render_credentials_email("Temp1234", "http://localhost:5173/login", "Portal de Calificaciones")
        assert "Temp1234" in test_html
        assert "Portal de Calificaciones" in test_html or "Tus credenciales" in test_html
        # Check current SMTP config (may be set in DB)
        real_cfg = _get_smtp_config()
        print(f"[MAIL] real_cfg host='{real_cfg.get('host')}' user='{real_cfg.get('user')}' (empty=>MOCK)")
        # Mock test: force empty host to verify MOCK path returns success without SMTP
        orig_get = mail_mod._get_smtp_config
        try:
            mail_mod._get_smtp_config = lambda: {"host": "", "port": 587, "user": "", "password": "", "use_tls": True, "sender": ""}
            res_mock = send_credentials_email("test@example.com", "Temp1234", "Alumno Test", "http://localhost:5173/login")
            assert res_mock.get("success") == True, f"mail mock failed {res_mock}"
            print(f"[MAIL] mock (forced empty) verified success={res_mock.get('success')} html_len={len(test_html)}")
        finally:
            mail_mod._get_smtp_config = orig_get
        # Also test real config path (if SMTP configured, may succeed or fail depending on creds; never log plaintext)
        res_real = send_credentials_email("test@example.com", "Temp1234", "Alumno Test", "http://localhost:5173/login")
        print(f"[MAIL] real send result success={res_real.get('success')} error={res_real.get('error','')} (MOCK if host empty, else SMTP attempt)")
        res = res_mock  # for summary flag

        # Summary
        print("="*60)
        print("SUMMARY")
        print(f"total_csv={total_csv} total_docx={total_docx} matched={matched} solo_boleta={len(only_boleta)} generic_emails={generic_emails}")
        print(f"carreras_normalizadas={dict(carrera_counter)} canonical_carreras={list(canon.keys())}")
        print(f"calificaciones_importadas={califs_imported} materias_nuevas={materias_new} alumnos_creados={created} alumnos_actualizados={updated} backup={backup_path} generic_csv={out_csv}")
        print(f"mail_mock={'OK' if res.get('success') else 'FAIL'} mode={'MOCK' if not os.environ.get('MAIL_SERVER') else 'SMTP'}")

if __name__ == "__main__":
    main()
