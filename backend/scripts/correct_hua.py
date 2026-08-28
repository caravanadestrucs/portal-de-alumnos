"""
Correct HUA/TEO assignment using UNIDAD field inside each DOCX (primary source).

- Parses every DOCX under project_root/boletas (86 files) with python-docx,
  extracts UNIDAD cell (HUAUTLA DE JIMENEZ vs TEOTITLAN DE FLORES MAGON)
  and every NOMBRE: entry (1 per individual DOCX, many per bulk DOCX).
- Builds mapping: normalized NOMBRE -> (sede determined by UNIDAD, source path).
  Bulk files have a single UNIDAD header that applies to all contained NOMBRE rows.
- For each alumno, resolution order:
  1) Exact normalized NOMBRE match (full and inverted) against parsed NOMBREs.
     If found, assign sede from that DOCX's UNIDAD (HUAUTLA->HUA, TEOTITLAN->TEO).
  2) Token-set fuzzy match (order-agnostic, Levenshtein per token <=2) and
     full-string Levenshtein <=2 against parsed NOMBREs. Same UNIDAD assignment.
  3) Fallback filename match (normalized stem without PENDIENTE/BAJA/TEMPORAL)
     against indexed DOCX filenames; sede then from that file's UNIDAD.
     If UNIDAD missing (none of the 86), fallback to folder name heuristic
     (huautla in path -> HUA, teotitlan -> TEO) — only for UNIDAD-less files.
  4) No evidence at all -> fallback TEO flagged for manual review.
- Updates Alumno.sede_id and rewrites backend/instance/manual_review.csv
  with only truly unmatched rows.

File-level UNIDAD counts (ground truth, not folder): 56 HUA, 30 TEO (86 total).
Bulk HUA files (6) contribute 41 distinct NOMBRE entries, all HUA by UNIDAD.

Usage:
  python backend/scripts/correct_hua.py
  python backend/scripts/correct_hua.py --dry-run
"""

import csv
import pathlib
import re
import sys
import unicodedata

BACKEND_DIR = pathlib.Path(__file__).resolve().parent.parent
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

PROJECT_ROOT = BACKEND_DIR.parent
BOLETAS_ROOT = PROJECT_ROOT / "boletas"

if not BOLETAS_ROOT.exists():
    alt = pathlib.Path(r"C:\Users\Dario\Desktop\portal de alumnos\boletas")
    if alt.exists():
        BOLETAS_ROOT = alt


def strip_accents(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFD", s) if unicodedata.category(c) != "Mn")


def normalize(s: str) -> str:
    """Lower, strip accents, keep alphanum, collapse spaces."""
    if not s:
        return ""
    s = strip_accents(s.lower())
    s = re.sub(r"[^a-z0-9]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def normalize_tokens(s: str):
    return normalize(s).split()


def lev(a: str, b: str) -> int:
    if a == b:
        return 0
    if len(a) < len(b):
        a, b = b, a
    prev = list(range(len(b) + 1))
    for i, ca in enumerate(a, 1):
        cur = [i]
        for j, cb in enumerate(b, 1):
            cur.append(min(cur[j - 1] + 1, prev[j] + 1, prev[j - 1] + (ca != cb)))
        prev = cur
    return prev[-1]


def tokens_match_set(a_tokens, b_tokens, max_lev=2) -> bool:
    """Order-agnostic token set match with per-token Levenshtein tolerance."""
    a_set = set(a_tokens)
    b_set = set(b_tokens)
    if len(a_set) != len(b_set):
        return False
    b_list = list(b_set)
    used = [False] * len(b_list)
    for at in a_set:
        found = False
        for j, bt in enumerate(b_list):
            if used[j]:
                continue
            if at == bt or lev(at, bt) <= max_lev:
                used[j] = True
                found = True
                break
        if not found:
            return False
    return True


# ---------------------------------------------------------------------------
# DOCX parsing helpers: UNIDAD is primary source, folder is fallback only
# ---------------------------------------------------------------------------

def extract_unidad_text(doc) -> str | None:
    """Return raw UNIDAD cell text (e.g. 'UNIDAD: HUAUTLA DE JIMENEZ, OAXACA') or None."""
    # Search tables first (boletas store UNIDAD in first table first row)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if "UNIDAD" in txt.upper():
                    # Return first line containing UNIDAD
                    for line in txt.splitlines():
                        if "UNIDAD" in line.upper():
                            return line.strip()
                    return txt.splitlines()[0].strip() if txt.splitlines() else txt.strip()
                for para in cell.paragraphs:
                    if "UNIDAD" in para.text.upper():
                        return para.text.strip()
    # Fallback search all paragraphs
    for para in doc.paragraphs:
        if "UNIDAD" in para.text.upper():
            return para.text.strip()
    return None


def determine_sede_from_unidad(unidad_text: str | None, fallback_path: pathlib.Path | None = None) -> str | None:
    """
    Determine sede code from UNIDAD text.
    Returns HUA/TEO/None. If unidad_text is None and fallback_path is given,
    uses folder name heuristic only as fallback (huautla/teotitlan in path).
    """
    if unidad_text:
        up = unidad_text.upper()
        if "HUAUTLA" in up:
            return "HUA"
        if "TEOTITLAN" in up:
            return "TEO"
        # Handle typo "HUAUTLA E JIMENEZ" etc is already caught by HUAUTLA
        # Unknown unidad string: fall through to fallback
    if fallback_path is not None:
        lower = str(fallback_path).lower()
        if "huautla" in lower:
            return "HUA"
        if "teotitlan" in lower:
            return "TEO"
    return None


def extract_nombres(doc) -> list[str]:
    """Collect every NOMBRE: value from tables."""
    nombres: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt.startswith("NOMBRE:"):
                    name = txt.replace("NOMBRE:", "").strip()
                    name = re.sub(r"\s+", " ", name).strip()
                    if name:
                        nombres.append(name)
    return nombres


def build_docx_and_unidad_maps():
    """
    Index all DOCX files and derive UNIDAD per file + NOMBRE->sede mapping.
    Returns:
      all_docx, file_unidad_map, nombre_to_sede, nombre_to_path, nombre_to_orig,
      docx_index, token_index, unidad_counts
    """
    from docx import Document

    all_docx = [p for p in BOLETAS_ROOT.rglob("*.docx") if not p.name.startswith("~$") and not p.name.startswith("~")]

    file_unidad_map: dict[str, str] = {}
    file_unidad_raw: dict[str, str] = {}
    nombre_to_sede: dict[str, str] = {}
    nombre_to_path: dict[str, str] = {}
    nombre_to_orig: dict[str, str] = {}
    nombre_to_tokens: dict[str, list[str]] = {}

    unidad_counts = {"HUA": 0, "TEO": 0, "UNK": 0}

    # For filename fallback (similar to previous script)
    docx_index: dict[str, list[pathlib.Path]] = {}
    token_index: dict[str, tuple[list[str], pathlib.Path, str]] = {}

    for p in all_docx:
        # Build filename indexes (for fallback path)
        cleaned = re.sub(r"(PENDIENTE|BAJA|TEMPORAL)", " ", p.stem, flags=re.I)
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        norm_key = normalize(cleaned)
        if norm_key:
            docx_index.setdefault(norm_key, []).append(p)
            orig_key = normalize(p.stem)
            if orig_key != norm_key:
                docx_index.setdefault(orig_key, []).append(p)
            tok = normalize_tokens(cleaned)
            token_index[norm_key] = (tok, p, p.stem)

        # Parse DOCX for UNIDAD and NOMBREs
        try:
            doc = Document(str(p))
        except Exception as e:
            print(f"[WARN] could not open {p}: {e}")
            # fallback to folder heuristic if unreadable
            sede_fb = determine_sede_from_unidad(None, p)
            sede = sede_fb or "UNK"
            file_unidad_map[str(p)] = sede
            unidad_counts[sede if sede in unidad_counts else "UNK"] += 1
            continue

        raw_unidad = extract_unidad_text(doc)
        sede = determine_sede_from_unidad(raw_unidad, p)
        if sede is None:
            sede = "UNK"
        file_unidad_map[str(p)] = sede
        file_unidad_raw[str(p)] = raw_unidad or ""
        unidad_counts[sede if sede in unidad_counts else "UNK"] += 1

        nombres = extract_nombres(doc)
        for name in nombres:
            norm_name = normalize(name)
            if not norm_name:
                continue
            if norm_name not in nombre_to_sede:
                nombre_to_sede[norm_name] = sede
                nombre_to_path[norm_name] = str(p)
                nombre_to_orig[norm_name] = name
                nombre_to_tokens[norm_name] = normalize_tokens(name)

    return (
        all_docx,
        file_unidad_map,
        file_unidad_raw,
        nombre_to_sede,
        nombre_to_path,
        nombre_to_orig,
        nombre_to_tokens,
        docx_index,
        token_index,
        unidad_counts,
    )


def main(dry_run=False):
    from app import create_app
    from models import db, Alumno, Sede

    app = create_app()
    with app.app_context():
        db.create_all()
        teo = Sede.query.filter_by(codigo="TEO").first()
        hua = Sede.query.filter_by(codigo="HUA").first()
        if not teo or not hua:
            from scripts.seed_sedes import seed_sedes
            seed_sedes()
            teo = Sede.query.filter_by(codigo="TEO").first()
            hua = Sede.query.filter_by(codigo="HUA").first()

        sede_map = {s.codigo: s.id for s in Sede.query.all()}
        if "TEO" not in sede_map or "HUA" not in sede_map:
            raise RuntimeError("Sede seeding failed")

        alumnos = Alumno.query.all()
        total = len(alumnos)
        print(f"[correct_hua] total alumnos {total}")
        print(f"[correct_hua] boletas root {BOLETAS_ROOT} exists={BOLETAS_ROOT.exists()}")

        (
            all_docx,
            file_unidad_map,
            file_unidad_raw,
            nombre_to_sede,
            nombre_to_path,
            nombre_to_orig,
            nombre_to_tokens,
            docx_index,
            token_index,
            unidad_counts,
        ) = build_docx_and_unidad_maps()

        print(f"[correct_hua] indexed DOCX {len(all_docx)} (86 expected)")
        print(f"[correct_hua] UNIDAD counts: HUA {unidad_counts['HUA']} TEO {unidad_counts['TEO']} UNK {unidad_counts['UNK']} (expected 56 HUA / 30 TEO)")
        hua_path_count = sum(1 for p in all_docx if "huautla" in str(p).lower())
        print(f"[correct_hua] folder 'huautla' path count {hua_path_count} (misleading, UNIDAD is truth; 56 HUA folders are mostly flat '2023/2024' without huautla)")
        # Log bulk files for visibility
        bulk_dir = BOLETAS_ROOT / "BOLETAS HUAUTLA 1ER. CUAT. SEP-DIC 2025"
        if bulk_dir.exists():
            bulk_files = [p for p in bulk_dir.iterdir() if p.suffix == ".docx" and not p.name.startswith("~")]
        else:
            bulk_files = [p for p in all_docx if "boletas huautla" in str(p).lower()]
        print(f"[correct_hua] bulk HUA files {len(bulk_files)}")
        for f in sorted(bulk_files):
            print(f"  - {f.name} -> UNIDAD={file_unidad_raw.get(str(f), '')!r} sede={file_unidad_map.get(str(f))}")
        print(f"[correct_hua] distinct NOMBRE entries parsed {len(nombre_to_sede)} (90 HUA + 30 TEO expected; bulk contributes 41)")
        # Quick sanity
        hua_nombres = sum(1 for v in nombre_to_sede.values() if v == "HUA")
        teo_nombres = sum(1 for v in nombre_to_sede.values() if v == "TEO")
        print(f"[correct_hua] NOMBRE->sede via UNIDAD: HUA {hua_nombres} TEO {teo_nombres}")

        to_teo = 0
        to_hua = 0
        flagged = 0
        details = []
        flagged_rows = []

        for alumno in alumnos:
            full = f"{alumno.nombre or ''} {alumno.apellido_paterno or ''} {alumno.apellido_materno or ''}".strip()
            full = re.sub(r"\s+", " ", full)
            inv = f"{alumno.apellido_paterno or ''} {alumno.apellido_materno or ''} {alumno.nombre or ''}".strip()
            inv = re.sub(r"\s+", " ", inv)

            norm_full = normalize(full)
            norm_inv = normalize(inv)
            tok_full = normalize_tokens(full)

            found_path = None
            found_sede = None
            found_nombre = None
            reason = ""
            needs_review = False
            codigo = "TEO"
            matched_via = None  # "unidad_nombre_exact", "unidad_nombre_fuzzy", "unidad_filename"

            # 1) Exact NOMBRE match via UNIDAD (primary truth)
            for cand_norm, cand_label in ((norm_full, "full"), (norm_inv, "inv")):
                if cand_norm in nombre_to_sede:
                    found_sede = nombre_to_sede[cand_norm]
                    found_path = nombre_to_path[cand_norm]
                    found_nombre = nombre_to_orig[cand_norm]
                    matched_via = "unidad_nombre_exact"
                    reason = f"unidad:{found_sede}:nombre_exact:{found_nombre} <- {cand_label}"
                    break

            # 2) Fuzzy NOMBRE match via UNIDAD (token-set and Levenshtein)
            if found_sede is None:
                for b_norm, b_tok in nombre_to_tokens.items():
                    matched = False
                    bulk_name = nombre_to_orig[b_norm]
                    # token-set order-agnostic
                    if tokens_match_set(tok_full, b_tok, max_lev=2):
                        matched = True
                    elif abs(len(norm_full) - len(b_norm)) <= 2 and lev(norm_full, b_norm) <= 2:
                        matched = True
                    elif abs(len(norm_inv) - len(b_norm)) <= 2 and lev(norm_inv, b_norm) <= 2:
                        matched = True
                    if matched:
                        found_sede = nombre_to_sede[b_norm]
                        found_path = nombre_to_path[b_norm]
                        found_nombre = bulk_name
                        matched_via = "unidad_nombre_fuzzy"
                        reason = f"unidad:{found_sede}:nombre_fuzzy:{found_nombre}"
                        break

            # 3) Fallback filename match, then resolve sede via UNIDAD of that file (not folder)
            if found_sede is None:
                fname_matched_path = None
                fname_reason = ""
                # exact filename
                for cand_norm in (norm_full, norm_inv):
                    if cand_norm in docx_index:
                        candidates = docx_index[cand_norm]
                        chosen = candidates[0]
                        # if multiple, prefer the one whose UNIDAD matches expectation? just pick first
                        fname_matched_path = str(chosen)
                        fname_reason = f"filename_exact:{cand_norm}"
                        break
                if fname_matched_path is None:
                    for norm_key, (tok, path, stem) in token_index.items():
                        if tokens_match_set(tok_full, tok, max_lev=2):
                            fname_matched_path = str(path)
                            fname_reason = f"filename_tokens:{normalize(stem)}"
                            break
                if fname_matched_path is not None:
                    # Resolve sede from that file's UNIDAD (fallback to folder only if UNIDAD was UNK)
                    sede_from_file = file_unidad_map.get(fname_matched_path)
                    if sede_from_file in ("HUA", "TEO"):
                        found_sede = sede_from_file
                    else:
                        # UNK -> fallback to folder heuristic per spec
                        lower = fname_matched_path.lower()
                        if "huautla" in lower:
                            found_sede = "HUA"
                            fname_reason += "+folder_fallback_huautla"
                        elif "teotitlan" in lower:
                            found_sede = "TEO"
                            fname_reason += "+folder_fallback_teotitlan"
                        else:
                            found_sede = "TEO"
                            fname_reason += "+fallback_unk_TEO"
                    found_path = fname_matched_path
                    raw_u = file_unidad_raw.get(fname_matched_path, "")
                    matched_via = "unidad_filename"
                    reason = f"unidad:{found_sede}:filename:{fname_reason}:{raw_u or 'no_unidad_raw'}"

            # 4) No evidence -> flagged TEO
            if found_sede is None:
                codigo = "TEO"
                reason = "fallback:assigned_TEO_flagged:no_docx_match_no_unidad"
                needs_review = True
                found_path = ""
            else:
                codigo = found_sede
                needs_review = False
                if codigo not in ("HUA", "TEO"):
                    codigo = "TEO"
                    reason += "+coerced_TEO"
                    needs_review = True

            if codigo == "TEO":
                to_teo += 1
            else:
                to_hua += 1
            if needs_review:
                flagged += 1
                flagged_rows.append(
                    {
                        "id": alumno.id,
                        "numero_control": alumno.numero_control,
                        "nombre_completo": alumno.nombre_completo,
                        "email": alumno.email,
                        "inferred_sede": codigo,
                        "reason": reason,
                        "needs_review": True,
                        "docx_path": found_path or "",
                    }
                )
            details.append(
                {
                    "id": alumno.id,
                    "numero_control": alumno.numero_control,
                    "nombre_completo": alumno.nombre_completo,
                    "inferred_sede": codigo,
                    "reason": reason,
                    "needs_review": needs_review,
                    "docx_path": found_path or "",
                    "matched_via": matched_via or "fallback",
                }
            )

            if not dry_run:
                alumno.sede_id = sede_map[codigo]

        if not dry_run:
            db.session.commit()

        output_csv = BACKEND_DIR / "instance" / "manual_review.csv"
        output_csv.parent.mkdir(parents=True, exist_ok=True)
        with open(output_csv, "w", newline="", encoding="utf-8") as f:
            if flagged == 0:
                f.write("# CORRECTED HUA/TEO assignment via UNIDAD - no manual review required\n")
                f.write(f"# Total {total} -> TEO {to_teo} HUA {to_hua} flagged 0\n")
                f.write(f"# Source: {BOLETAS_ROOT} ({len(all_docx)} DOCX, UNIDAD HUA {unidad_counts['HUA']} TEO {unidad_counts['TEO']})\n")
                f.write(f"# NOMBRE entries {len(nombre_to_sede)} (HUA {hua_nombres} TEO {teo_nombres})\n")
            writer = csv.DictWriter(
                f,
                fieldnames=[
                    "id",
                    "numero_control",
                    "nombre_completo",
                    "email",
                    "inferred_sede",
                    "reason",
                    "needs_review",
                ],
            )
            writer.writeheader()
            for row in flagged_rows:
                writer.writerow({k: row[k] for k in ["id", "numero_control", "nombre_completo", "email", "inferred_sede", "reason", "needs_review"]})

        print(f"[correct_hua] dry_run={dry_run} total={total} TEO={to_teo} HUA={to_hua} flagged={flagged}")
        if dry_run:
            print("[correct_hua] dry-run: no DB writes")
        else:
            print(f"[correct_hua] DB updated, manual_review.csv written to {output_csv} flagged={flagged}")
        # Extra verification hint for task step 4
        print(f"[correct_hua] EXPECTED ~30 TEO / 73-79 HUA (via UNIDAD), not 79 TEO / 30 HUA (old folder bug).")

        return {
            "total": total,
            "to_teo": to_teo,
            "to_hua": to_hua,
            "flagged": flagged,
            "details": details,
            "flagged_rows": flagged_rows,
            "dry_run": dry_run,
            "unidad_counts": unidad_counts,
        }


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Correct HUA/TEO via UNIDAD field inside DOCX")
    parser.add_argument("--dry-run", action="store_true", help="Report only, no DB writes")
    args = parser.parse_args()
    report = main(dry_run=args.dry_run)
    print(f"{report['to_teo']} TEO, {report['to_hua']} HUA, flagged {report['flagged']}")
